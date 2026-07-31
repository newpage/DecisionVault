import io,json,os,time
from pathlib import Path
from uuid import uuid4
import httpx
from docx import Document
from pypdf import PdfReader
from sqlalchemy import create_engine,select
from sqlalchemy.orm import Session
from sqlalchemy.ext.automap import automap_base

DB=os.environ["DATABASE_URL"]
STORAGE=Path(os.getenv("STORAGE_PATH","/data/storage"))
OLLAMA=os.getenv("OLLAMA_URL","http://host.docker.internal:11434")
EMBED_MODEL=os.getenv("OLLAMA_EMBED_MODEL","nomic-embed-text")
engine=create_engine(DB,pool_pre_ping=True)
Base=automap_base()

def extract(name:str,raw:bytes)->str:
    lower=name.lower()
    if lower.endswith('.pdf'):
        if not raw.startswith(b'%PDF-'): raise ValueError('Invalid PDF signature')
        return '\n'.join(page.extract_text() or '' for page in PdfReader(io.BytesIO(raw)).pages)
    if lower.endswith('.docx'):
        doc=Document(io.BytesIO(raw)); parts=[p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows: parts.append(' | '.join(cell.text for cell in row.cells))
        return '\n'.join(parts)
    if lower.endswith(('.txt','.md','.csv','.json')): return raw.decode('utf-8',errors='replace')
    raise ValueError('Unsupported file type. Supported: PDF, DOCX, TXT, MD, CSV, JSON')

def chunks(text:str,size=1200,overlap=180):
    clean=' '.join(text.split()); out=[]; start=0
    while start<len(clean): out.append(clean[start:start+size]); start+=size-overlap
    return out or ['']

def embed(text:str):
    try:
        r=httpx.post(f'{OLLAMA}/api/embed',json={'model':EMBED_MODEL,'input':text},timeout=20);r.raise_for_status();return (r.json().get('embeddings') or [None])[0]
    except Exception:return None

def main():
    while True:
        try:
            Base.prepare(autoload_with=engine)
            Job=Base.classes.ingestion_jobs; Source=Base.classes.source_documents; Card=Base.classes.knowledge_cards; Evidence=Base.classes.knowledge_evidence; Chunk=Base.classes.knowledge_chunks; Audit=Base.classes.audit_events
            with Session(engine) as db:
                job=db.scalar(select(Job).where(Job.status=='queued').order_by(Job.created_at).with_for_update(skip_locked=True))
                if not job: time.sleep(2); continue
                job.status='processing';job.progress=10;db.commit()
                source=db.get(Source,job.source_document_id)
                try:
                    text=extract(source.filename,(STORAGE/source.storage_key).read_bytes())
                    if not text.strip(): raise ValueError('No extractable text found; OCR is intentionally not enabled')
                    card=Card(id=str(uuid4()),tenant_id=source.tenant_id,workspace_id=source.workspace_id,title=Path(source.filename).stem,summary=text[:320],body=text,knowledge_type='source_derived',lifecycle_status='draft',approval_status='not_submitted',authority_level='unverified',classification_rank=20,ai_usage_allowed=True,trust_score=.45,owner_id=source.created_by,created_at=job.created_at)
                    db.add(card);db.flush()
                    for i,part in enumerate(chunks(text)):
                        db.add(Chunk(id=str(uuid4()),tenant_id=source.tenant_id,knowledge_card_id=card.id,content=part,chunk_index=i,search_text=part.lower(),embedding=embed(part)))
                    db.add(Evidence(id=str(uuid4()),tenant_id=source.tenant_id,knowledge_card_id=card.id,source_document_id=source.id,locator='Extracted content',excerpt=text[:1000]))
                    db.add(Audit(id=str(uuid4()),tenant_id=source.tenant_id,actor_id=source.created_by,event_type='KnowledgeDraftCreated',entity_type='knowledge_card',entity_id=card.id,description=f'Draft Knowledge Card created from {source.filename}',details={}))
                    source.status='processed';job.status='completed';job.progress=100;db.commit()
                except Exception as exc:
                    source.status='failed';job.status='failed';job.error=str(exc);db.commit()
        except Exception as exc:
            print(f'worker loop error: {exc}',flush=True);time.sleep(5)
if __name__=='__main__': main()
